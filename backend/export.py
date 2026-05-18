import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATES = {
    "minimal":  "Minimal — Klar & übersichtlich",
    "academic": "Akademisch — Strukturiert & formal",
    "modern":   "Modern — Farbig & professionell",
}


def export_to_docx(title: str, content: str, template: str = "minimal") -> bytes:
    doc = Document()
    handlers = {"minimal": _minimal, "academic": _academic, "modern": _modern}
    handlers.get(template, _minimal)(doc, title, content)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── helpers ──────────────────────────────────────────────────────────────────

def _hr(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade(paragraph, hex_color: str):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _parse_content(doc, content: str, h1_size=14, h2_size=12,
                   h1_color=(30, 30, 30), h2_color=(60, 60, 60),
                   indent_cm=0.0):
    for line in content.split("\n"):
        if line.startswith("# "):
            p = doc.add_paragraph()
            r = p.add_run(line[2:])
            r.bold = True
            r.font.size = Pt(h1_size)
            r.font.color.rgb = RGBColor(*h1_color)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(4)
        elif line.startswith("## "):
            p = doc.add_paragraph()
            r = p.add_run(line[3:])
            r.bold = True
            r.font.size = Pt(h2_size)
            r.font.color.rgb = RGBColor(*h2_color)
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(4)
        elif line.strip():
            p = doc.add_paragraph(line.strip())
            p.paragraph_format.space_after = Pt(6)
            if indent_cm:
                p.paragraph_format.first_line_indent = Cm(indent_cm)


# ── templates ─────────────────────────────────────────────────────────────────

def _minimal(doc, title, content):
    s = doc.sections[0]
    s.left_margin = Cm(3)
    s.right_margin = Cm(3)

    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(20, 20, 20)
    p.paragraph_format.space_after = Pt(4)

    d = doc.add_paragraph()
    dr = d.add_run(datetime.now().strftime("%d.%m.%Y"))
    dr.font.size = Pt(9)
    dr.font.color.rgb = RGBColor(160, 160, 160)
    d.paragraph_format.space_after = Pt(18)

    _hr(doc)
    _parse_content(doc, content)


def _academic(doc, title, content):
    s = doc.sections[0]
    s.left_margin = Cm(3.5)
    s.right_margin = Cm(3.5)

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = label.add_run("STUDIENDOKUMENT")
    lr.font.size = Pt(8)
    lr.font.color.rgb = RGBColor(120, 120, 120)
    lr.font.all_caps = True
    label.paragraph_format.space_after = Pt(6)

    tp = doc.add_paragraph()
    tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tp.add_run(title)
    tr.bold = True
    tr.font.size = Pt(20)
    tp.paragraph_format.space_after = Pt(4)

    dp = doc.add_paragraph()
    dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dpr = dp.add_run(datetime.now().strftime("%d. %B %Y"))
    dpr.font.size = Pt(10)
    dpr.font.color.rgb = RGBColor(110, 110, 110)
    dp.paragraph_format.space_after = Pt(20)

    _hr(doc)
    _parse_content(doc, content, h1_size=14, h2_size=12,
                   h1_color=(30, 30, 30), h2_color=(70, 70, 70), indent_cm=0.5)


def _modern(doc, title, content):
    s = doc.sections[0]
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

    header = doc.add_paragraph()
    _shade(header, "4F46E5")
    hr = header.add_run(f"  {title}")
    hr.bold = True
    hr.font.size = Pt(22)
    hr.font.color.rgb = RGBColor(255, 255, 255)
    header.paragraph_format.space_after = Pt(0)

    sub = doc.add_paragraph()
    _shade(sub, "6366F1")
    sr = sub.add_run(f"  {datetime.now().strftime('%d.%m.%Y')}  ·  Study Assistant")
    sr.font.size = Pt(9)
    sr.font.color.rgb = RGBColor(220, 220, 255)
    sub.paragraph_format.space_after = Pt(18)

    _parse_content(doc, content, h1_size=15, h2_size=12,
                   h1_color=(79, 70, 229), h2_color=(99, 102, 241))
